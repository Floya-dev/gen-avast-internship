import os
import re
import zipfile
from datetime import datetime
from docx import Document
import xml.etree.ElementTree as ET
from pathlib import Path

class DocxScorer:
    def __init__(self, root_folder):
        self.root_folder = root_folder
        self.file_scores = []
        
    def get_core_metadata(self, docx_path):
        """Extract core metadata from docx file"""
        metadata = {
            "created": None,
            "modified": None,
        }

        try:
            with zipfile.ZipFile(docx_path) as z:
                if "docProps/core.xml" in z.namelist():
                    core_xml = z.read("docProps/core.xml")
                    root = ET.fromstring(core_xml)

                    ns = {
                        "dcterms": "http://purl.org/dc/terms/"
                    }

                    created = root.find("dcterms:created", ns)
                    modified = root.find("dcterms:modified", ns)

                    if created is not None:
                        metadata["created"] = created.text
                    if modified is not None:
                        metadata["modified"] = modified.text
        except Exception as e:
            pass

        return metadata

    def count_images(self, docx_path):
        """Count images in the docx file"""
        count = 0
        try:
            with zipfile.ZipFile(docx_path) as z:
                # Count image files in the media folder
                for filename in z.namelist():
                    if filename.startswith('word/media/'):
                        count += 1
        except:
            pass
        return count

    def count_links(self, doc):
        """Count hyperlinks in document"""
        count = 0
        try:
            rels = doc.part.rels
            for rel in rels:
                if "hyperlink" in rels[rel].reltype:
                    count += 1
        except:
            pass
        return count

    def analyze_headings(self, doc):
        """Analyze heading structure"""
        headings = []
        for p in doc.paragraphs:
            try:
                if p.style and p.style.name and p.style.name.startswith("Heading"):
                    headings.append(p.style.name)
            except:
                continue

        if not headings:
            return {"no_headings": True, "no_hierarchy": True}

        levels = []
        for h in headings:
            match = re.search(r"\d+", h)
            if match:
                levels.append(int(match.group()))

        no_hierarchy = len(set(levels)) <= 1
        return {
            "no_headings": False,
            "no_hierarchy": no_hierarchy
        }

    def parse_date(self, date_string):
        """Parse ISO format date string"""
        if not date_string:
            return None
        try:
            return datetime.fromisoformat(date_string.replace("Z", "+00:00"))
        except:
            return None

    def extract_text_from_doc(self, doc):
        """Extract all text from paragraphs and tables"""
        full_text = ""
        
        # Get text from paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                full_text += p.text + " "
        
        # Get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            full_text += p.text + " "
        
        return full_text

    def calculate_score(self, metadata, full_metadata):
        """Calculate score based on all scoring rules"""
        score = 0
        details = []
        
        # Author scoring
        author = metadata['author']
        if author == 'BLANK':
            score += 1
            details.append(f"author=BLANK (+1)")
        elif author in ['python-docx', 'html-to-docx']:
            score += 2
            details.append(f"author={author} (+2)")
        else:
            score -= 1
            details.append(f"author={author} (-1)")
        
        # Keywords scoring
        keywords = metadata['keywords']
        if keywords == 'html-to-docx':
            score += 1
            details.append(f"keywords=html-to-docx (+1)")
        
        # Last modified by scoring
        last_modified_by = metadata['last_modified_by']
        if last_modified_by == 'BLANK':
            score += 1
            details.append(f"last_modified_by=BLANK (+1)")
        elif last_modified_by in ['python-docx', 'html-to-docx']:
            score += 2
            details.append(f"last_modified_by={last_modified_by} (+2)")
        elif last_modified_by == 'User':
            score += 1
            details.append(f"last_modified_by={last_modified_by} (+1)")
        else:
            score -= 2
            details.append(f"last_modified_by={last_modified_by} (-2)")
        
        # Revision scoring
        revision = metadata['revision']
        if revision == 'BLANK':
            score += 2
            details.append(f"revision=BLANK (+2)")
        elif revision == '1':
            score += 1
            details.append(f"revision=1 (+1)")
        else:
            score -= 2
            details.append(f"revision={revision} (-2)")
        
        # Links scoring
        links = full_metadata['link_count']
        if links > 0:
            score += 1
            details.append(f"links={links} (+1)")
        else:
            score -= 1
            details.append(f"links=0 (-1)")
        
        # Image scoring
        image_count = full_metadata['image_count']
        if image_count == 2:
            score += 2
            details.append(f"images=2 (+2)")
        
        # Date comparison
        created = full_metadata['created_date']
        modified = full_metadata['modified_date']
        
        if created and modified:
            if created == modified:
                score += 1
                details.append(f"creation_date=modified_date (+1)")
            elif created > modified:
                score += 3
                details.append(f"creation_date > modified_date (+3)")
        
        # Heading analysis
        heading_info = full_metadata['heading_info']
        
        if heading_info["no_headings"]:
            score += 1
            details.append(f"no_headings (+1)")
        
        if heading_info["no_hierarchy"]:
            score += 1
            details.append(f"no_hierarchy (+1)")
        
        return score, details
    
    def extract_metadata(self, docx_path, doc):
        """Extract relevant metadata from a .docx file"""
        try:
            core_props = doc.core_properties
            
            metadata = {
                'author': core_props.author.strip() if core_props.author else 'BLANK',
                'keywords': core_props.keywords.strip() if core_props.keywords else 'BLANK',
                'last_modified_by': core_props.last_modified_by.strip() if core_props.last_modified_by else 'BLANK',
                'revision': str(core_props.revision) if core_props.revision else 'BLANK',
            }
            
            return metadata
        except Exception as e:
            print(f"Error reading {docx_path}: {str(e)}")
            return None
    
    def score_files(self, show_details=False):
        """Score all .docx files in the folder"""
        print(f"Scoring files in: {self.root_folder}\n")
        
        for root, dirs, files in os.walk(self.root_folder):
            for file in files:
                if file.endswith('.docx') and not file.startswith('~$'):
                    file_path = os.path.join(root, file)
                    
                    try:
                        doc = Document(file_path)
                        metadata = self.extract_metadata(file_path, doc)
                        
                        if metadata:
                            # Get core dates
                            core_metadata = self.get_core_metadata(file_path)
                            created = self.parse_date(core_metadata["created"])
                            modified = self.parse_date(core_metadata["modified"])
                            
                            # Count links
                            link_count = self.count_links(doc)
                            
                            # Count images
                            image_count = self.count_images(file_path)
                            
                            # Analyze headings
                            heading_info = self.analyze_headings(doc)
                            
                            # Extract text (for future use if needed)
                            full_text = self.extract_text_from_doc(doc)
                            
                            words = re.findall(r"\b\w+\b", full_text.lower())
                            total_words = len(words) if words else 0
                            
                            # Full metadata for scoring
                            full_metadata = {
                                'created_date': created,
                                'modified_date': modified,
                                'link_count': link_count,
                                'image_count': image_count,
                                'heading_info': heading_info,
                                'total_words': total_words
                            }
                            
                            score, details = self.calculate_score(metadata, full_metadata)
                            
                            relative_path = os.path.relpath(file_path, self.root_folder)
                            
                            self.file_scores.append({
                                'path': file_path,
                                'relative_path': relative_path,
                                'filename': file,
                                'score': score,
                                'details': details,
                                'metadata': metadata,
                                'full_metadata': full_metadata
                            })
                            
                            if show_details:
                                print(f"Processed: {relative_path} - Score: {score}")
                    except Exception as e:
                        print(f"Error processing {file_path}: {str(e)}")
    
    def print_results(self, sort_by='score', show_details=False):
        """Print scoring results"""
        if not self.file_scores:
            print("No files were scored!")
            return
        
        # Sort files
        if sort_by == 'score':
            sorted_files = sorted(self.file_scores, key=lambda x: (-x['score'], x['filename']))
        else:  # sort by name
            sorted_files = sorted(self.file_scores, key=lambda x: x['filename'])
        
        print("\n" + "="*120)
        print("FILE SCORING RESULTS")
        print("="*120)
        print(f"Total files scored: {len(self.file_scores)}\n")
        
        # Calculate statistics
        scores = [f['score'] for f in self.file_scores]
        avg_score = sum(scores) / len(scores) if scores else 0
        max_score = max(scores) if scores else 0
        min_score = min(scores) if scores else 0
        
        print(f"Average Score: {avg_score:.2f}")
        print(f"Highest Score: {max_score}")
        print(f"Lowest Score:  {min_score}")
        print(f"Score Range:   {min_score} to {max_score}")
        
        # Score distribution
        print("\nScore Distribution:")
        score_dist = {}
        for f in self.file_scores:
            score_dist[f['score']] = score_dist.get(f['score'], 0) + 1
        
        for score in sorted(score_dist.keys(), reverse=True):
            count = score_dist[score]
            bar = '█' * count
            print(f"  Score {score:+3d}: {count:4d} files {bar}")
        
        print("\n" + "="*120)
        print(f"{'SCORE':<8} {'FILE':<75} {'PATH':<35}")
        print("-"*120)
        
        for item in sorted_files:
            score_str = f"{item['score']:+3d}"
            filename = item['filename'][:73]
            path = os.path.dirname(item['relative_path'])[:33]
            
            print(f"{score_str:<8} {filename:<75} {path:<35}")
            
            if show_details:
                for detail in item['details']:
                    print(f"         └─ {detail}")
                print()
    
    def print_detailed_report(self):
        """Print detailed report with metadata breakdown"""
        if not self.file_scores:
            print("No files were scored!")
            return
        
        sorted_files = sorted(self.file_scores, key=lambda x: (-x['score'], x['filename']))
        
        print("\n" + "="*120)
        print("DETAILED SCORING REPORT")
        print("="*120)
        
        for idx, item in enumerate(sorted_files, 1):
            print(f"\n[{idx}] {item['filename']}")
            print(f"    Path: {item['relative_path']}")
            print(f"    SCORE: {item['score']:+d}")
            print(f"    Breakdown:")
            for detail in item['details']:
                print(f"      • {detail}")
            print(f"    Metadata:")
            for key, value in item['metadata'].items():
                print(f"      - {key}: {value}")
            print(f"    Document Analysis:")
            print(f"      - Links: {item['full_metadata']['link_count']}")
            print(f"      - Images: {item['full_metadata']['image_count']}")
            print(f"      - Total Words: {item['full_metadata']['total_words']}")
            print(f"      - Has Headings: {not item['full_metadata']['heading_info']['no_headings']}")
            print(f"      - Has Hierarchy: {not item['full_metadata']['heading_info']['no_hierarchy']}")
            if item['full_metadata']['created_date']:
                print(f"      - Created: {item['full_metadata']['created_date']}")
            if item['full_metadata']['modified_date']:
                print(f"      - Modified: {item['full_metadata']['modified_date']}")
    
    def export_to_csv(self, output_file='file_scores.csv', include_metadata=False):
        """Export scores to CSV file"""
        import csv
        
        with open(output_file, 'w', newline='', encoding='utf-8') as f:
            if include_metadata:
                fieldnames = ['score', 'filename', 'relative_path', 'author', 'keywords', 
                            'last_modified_by', 'revision', 'links', 'images', 'total_words',
                            'has_headings', 'has_hierarchy', 'scoring_details']
            else:
                fieldnames = ['score', 'filename', 'relative_path']
            
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            
            sorted_files = sorted(self.file_scores, key=lambda x: (-x['score'], x['filename']))
            
            for item in sorted_files:
                row = {
                    'score': item['score'],
                    'filename': item['filename'],
                    'relative_path': item['relative_path']
                }
                
                if include_metadata:
                    row.update({
                        'author': item['metadata']['author'],
                        'keywords': item['metadata']['keywords'],
                        'last_modified_by': item['metadata']['last_modified_by'],
                        'revision': item['metadata']['revision'],
                        'links': item['full_metadata']['link_count'],
                        'images': item['full_metadata']['image_count'],
                        'total_words': item['full_metadata']['total_words'],
                        'has_headings': not item['full_metadata']['heading_info']['no_headings'],
                        'has_hierarchy': not item['full_metadata']['heading_info']['no_hierarchy'],
                        'scoring_details': '; '.join(item['details'])
                    })
                
                writer.writerow(row)
        
        print(f"\n✓ Scores exported to: {output_file}")
    
    def export_to_json(self, output_file='file_scores.json'):
        """Export scores to JSON file"""
        import json
        
        sorted_files = sorted(self.file_scores, key=lambda x: (-x['score'], x['filename']))
        
        # Calculate statistics
        scores = [f['score'] for f in self.file_scores]
        
        results = {
            'summary': {
                'total_files': len(self.file_scores),
                'average_score': sum(scores) / len(scores) if scores else 0,
                'max_score': max(scores) if scores else 0,
                'min_score': min(scores) if scores else 0
            },
            'files': [
                {
                    'filename': item['filename'],
                    'relative_path': item['relative_path'],
                    'score': item['score'],
                    'details': item['details'],
                    'metadata': item['metadata'],
                    'document_analysis': {
                        'links': item['full_metadata']['link_count'],
                        'images': item['full_metadata']['image_count'],
                        'total_words': item['full_metadata']['total_words'],
                        'has_headings': not item['full_metadata']['heading_info']['no_headings'],
                        'has_hierarchy': not item['full_metadata']['heading_info']['no_hierarchy'],
                        'created_date': str(item['full_metadata']['created_date']) if item['full_metadata']['created_date'] else None,
                        'modified_date': str(item['full_metadata']['modified_date']) if item['full_metadata']['modified_date'] else None
                    }
                }
                for item in sorted_files
            ]
        }
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False)
        
        print(f"✓ Scores exported to: {output_file}")
    
    def get_files_by_score_range(self, min_score=None, max_score=None):
        """Get files within a specific score range"""
        filtered = self.file_scores
        
        if min_score is not None:
            filtered = [f for f in filtered if f['score'] >= min_score]
        
        if max_score is not None:
            filtered = [f for f in filtered if f['score'] <= max_score]
        
        return sorted(filtered, key=lambda x: (-x['score'], x['filename']))


def main():
    print("="*120)
    print("DOCX FILE SCORER")
    print("="*120)
    print("\nScoring Rules:")
    print("  Author: BLANK=+1, python-docx/html-to-docx=+2, Other=-1")
    print("  Keywords: html-to-docx=+1")
    print("  Last Modified By: BLANK=+1, python-docx/html-to-docx=+2, Other=-2")
    print("  Revision: BLANK=+2, 1=+1, Other=-2")
    print("  Links: Found=+1, Not Found=-1")
    print("  Images: Exactly 2 images=+2")
    print("  Dates: creation_date=modified_date=+1, creation_date>modified_date=+3")
    print("  Headings: No headings=+1, No hierarchy=+1")
    print("="*120)
    
    # Get folder path
    FOLDER_PATH = input("\nEnter the folder path to analyze (or press Enter for current directory): ").strip()
    
    if not FOLDER_PATH:
        FOLDER_PATH = '.'
    
    if not os.path.exists(FOLDER_PATH):
        print(f"Error: Folder '{FOLDER_PATH}' does not exist!")
        return
    
    # Initialize scorer
    scorer = DocxScorer(FOLDER_PATH)
    
    # Score files
    print("\nProcessing files...")
    scorer.score_files(show_details=False)
    
    if not scorer.file_scores:
        print("\nNo .docx files found in the specified folder!")
        return
    
    # Display options
    while True:
        print("\n" + "="*120)
        print("DISPLAY OPTIONS:")
        print("  1. Summary view (sorted by score)")
        print("  2. Summary view with details")
        print("  3. Detailed report (full breakdown)")
        print("  4. Filter by score range")
        print("  5. Export to CSV")
        print("  6. Export to JSON")
        print("  7. Exit")
        print("="*120)
        
        choice = input("\nSelect option (1-7): ").strip()
        
        if choice == '1':
            scorer.print_results(sort_by='score', show_details=False)
        
        elif choice == '2':
            scorer.print_results(sort_by='score', show_details=True)
        
        elif choice == '3':
            scorer.print_detailed_report()
        
        elif choice == '4':
            try:
                min_score = input("Enter minimum score (or press Enter for no limit): ").strip()
                max_score = input("Enter maximum score (or press Enter for no limit): ").strip()
                
                min_score = int(min_score) if min_score else None
                max_score = int(max_score) if max_score else None
                
                filtered = scorer.get_files_by_score_range(min_score, max_score)
                
                print(f"\nFiles with score range {min_score or '-∞'} to {max_score or '+∞'}:")
                print("-"*120)
                for item in filtered:
                    print(f"[{item['score']:+3d}] {item['relative_path']}")
                print(f"\nTotal: {len(filtered)} files")
            except ValueError:
                print("Invalid input! Please enter numeric values.")
        
        elif choice == '5':
            include_meta = input("Include metadata in CSV? (y/n): ").strip().lower() == 'y'
            scorer.export_to_csv(include_metadata=include_meta)
        
        elif choice == '6':
            scorer.export_to_json()
        
        elif choice == '7':
            print("\nGoodbye!")
            break
        
        else:
            print("Invalid option! Please select 1-7.")


if __name__ == "__main__":
    main()
